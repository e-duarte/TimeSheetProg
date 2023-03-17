from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Pt

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(.5)
        section.bottom_margin = Cm(.5)
        section.left_margin = Cm(1.)
        section.right_margin = Cm(1.)

def set_doc_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

def set_column_width(cells, width):
    for cell in cells:
        cell.width = width

def set_assign_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r in table.rows:
        for c in r.cells:
            # c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            

def set_color_cell(cell, color='5b9bd5'):
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)

def set_header_doc_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
    for r in table.rows:
        for c in r.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
    set_column_width(table.columns[0].cells, Cm(5.))
    set_column_width(table.columns[1].cells, Cm(10.))
    set_column_width(table.columns[2].cells, Cm(5.))

def set_table_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    par = table.cell(0,0).paragraphs[-1]
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_color_cell(table.cell(0,0))
    
    table.style = 'Table Grid'
    table.allow_autofit = False

    columns = table.columns
    
    set_column_width(columns[0].cells, Cm(1.))

    aux_col = [c for i, c in enumerate(columns) if i != 0]
    for i in range(0, len(aux_col), 5):
        set_column_width(aux_col[i].cells, Cm(.5))
        set_color_cell(table.cell(2,i+1), '70ad47')
        
        set_column_width(aux_col[i+1].cells, Cm(.5))
        set_column_width(aux_col[i+2].cells, Cm(.5))
        set_color_cell(table.cell(2,i+2+1), 'ffc000')
        set_column_width(aux_col[i+3].cells, Cm(.5))
        set_column_width(aux_col[i+4].cells, Cm(5.8) if len(aux_col) == 10 else Cm(15.))
        table.cell(2,i+4+1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in table.rows[2:-1]:
        for c in row.cells:
            c.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


    