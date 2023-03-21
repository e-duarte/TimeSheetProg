from docx.shared import Cm, Pt
from calendar_util import CalendarUtil
from docx.enum.table import WD_ROW_HEIGHT_RULE
from styles import set_table_style, set_margins, set_doc_style
from datetime import datetime



class TimeSheet:
    def __init__(self, document, institution, title, employee, month):
        self.institution = institution
        self.title = title
        self.employee = employee
        self.document = document
        self.month = month

        self.table = self.build_table()
        self.build_table_header()
        self.build_body()
        self.build_footer()

    def build_table(self):
        calendar = CalendarUtil(month=self.month)

        header_rows_num = 2
        header_values_row_num = 2
        body_row_num = len(calendar.get_current_month_range())
        footer_row_num = 1
        rows_num = header_rows_num + header_values_row_num + body_row_num + footer_row_num

        days_col_num = 1
        turn_col_num = 2
        fault_col_num = 2
        signature_col_num = 1
        columns_num = days_col_num + ((turn_col_num + fault_col_num + signature_col_num) if len(self.employee['Turnos']) == 1 else (turn_col_num + fault_col_num + signature_col_num)*2)

        table = self.document.add_table(rows=rows_num, cols=columns_num)
        set_table_style(table)

        return table
    
    def add_header_text(self):
        start = self.table.cell(2, 0)
        end = self.table.cell(3, 0)
        date_cell = start.merge(end)
        run = date_cell.paragraphs[-1].add_run('Data')
        run.font.bold = True
        run.font.size = Pt(10)
        
        cell_index = 1
        for turn in self.employee['Turnos']:
            start = self.table.cell(2, cell_index)
            end = self.table.cell(2, cell_index+1)

            turn_cell = start.merge(end)
            run = turn_cell.paragraphs[-1].add_run(f'Turno {turn}º')
            run.font.bold = True
            run.font.size = Pt(10)

            in_cell = self.table.cell(3, cell_index)
            run = in_cell.paragraphs[-1].add_run(f'Entrada')
            run.font.bold = True
            run.font.size = Pt(10)
            out_cell = self.table.cell(3, cell_index+1)
            run = out_cell.paragraphs[-1].add_run(f'Saída')
            run.font.bold = True
            run.font.size = Pt(10)

            cell_index += 2

            start = self.table.cell(2, cell_index)
            end = self.table.cell(2, cell_index+1)

            fault_cell = start.merge(end)
            run = fault_cell.paragraphs[-1].add_run('Faltas')
            run.font.bold = True
            run.font.size = Pt(10)

            in_cell = self.table.cell(3, cell_index)
            run = in_cell.paragraphs[-1].add_run(f'Just.')
            run.font.bold = True
            run.font.size = Pt(10)
            out_cell = self.table.cell(3, cell_index+1)
            run = out_cell.paragraphs[-1].add_run(f'Ñ Just.')
            run.font.bold = True
            run.font.size = Pt(10)

            cell_index += 2
            start = self.table.cell(2, cell_index)
            end = self.table.cell(3, cell_index)
            assign_cell = start.merge(end)

            run = assign_cell.paragraphs[-1].add_run('Assinatura do servidor')
            run.font.bold = True
            run.font.size = Pt(10)

            cell_index += 1
        
    def add_line_metadata(self, par, line_data):
        formated_data = []
        for i, data in enumerate(line_data):
            if i == 0:
                formated_data.append((data[0], data[1]))
            else:
                formated_data.append((f'    {data[0]}', data[1]))

        for data in formated_data:
            #field
            run = par.add_run(f'{data[0]}')
            run.font.bold = True
            #value
            run = par.add_run(f'{data[1]}')
        
    def build_table_header(self):
        calendar = CalendarUtil(month=self.month)
        none_value = '_________________'

        title_row = self.table.rows[0]
        title_row = title_row.cells[0].merge(title_row.cells[-1])
        title_run = title_row.paragraphs[-1].add_run(self.title)
        title_run.font.bold = True

        header_row = self.table.rows[1]
        header_row = header_row.cells[0].merge(header_row.cells[-1])

        self.add_line_metadata(header_row.paragraphs[-1], ((self.institution, ''),))
        self.add_line_metadata(
            header_row.paragraphs[-1],
            (
                ('\nNome do (a) Servidor (a): ', f'{self.employee["Nome"]}'),
                ('Carga Horária: ', f'{self.employee["Carga Horária"]}h'),
            )
        )
        self.add_line_metadata(
            header_row.paragraphs[-1],
            (
                ('\nMatrícula: ', none_value if self.employee['Matrícula'] == '' else self.employee['Matrícula']),
                ('Mês/Ano: ', f'{calendar.months[self.month-1]}/{datetime.today().year}'),
                ('Situação Funcional: ', f'{self.employee["Situação"]}')
            )
        )

        self.add_line_metadata(
            header_row.paragraphs[-1],
            (
                ('\nCargo/Função: ', f'{self.employee["Função"]}'),
                ('Telefone: ', f'(93){none_value}' if self.employee['Telefone'] == '' else f'(93) {self.employee["Telefone"]}'),
            )
        )    

        self.add_header_text()

    def build_body(self):
        days = CalendarUtil(month=self.month).get_current_month_range()

        for day, row in zip(days, self.table.rows[4:-1]):
            row.cells[0].paragraphs[-1].text = str(day)

            weekday = CalendarUtil(month=self.month).get_weekday(day)
            if weekday[1] == 'Sábado' or weekday[1] == 'Domingo':
                cells = row.cells[1:]
                cell_index = 0
                for i in range(len(self.employee['Turnos'])):
                    cells[cell_index].paragraphs[-1].text = 'XXX'
                    cell_index += 1
                    cells[cell_index].paragraphs[-1].text = 'XXX'
                    cell_index += 1
                    cells[cell_index].paragraphs[-1].text = 'XXX'
                    cell_index += 1
                    cells[cell_index].paragraphs[-1].text = 'XXX'
                    cell_index += 1
                    cells[cell_index].paragraphs[-1].text = weekday[1]
                    cell_index = 5
                    
    def build_footer(self):
        start = self.table.rows[-1].cells[0]
        end = self.table.rows[-1].cells[-1]

        merged_cell = start.merge(end)
        merged_cell.paragraphs[-1].add_run(
            'OBS.: Anexar motivo da ausência e/ou atestado médico, quando aplicável'
        )

if __name__ == '__main__':
    from docx import Document
    from main import EmployeesLoader

    e = EmployeesLoader('funcionarios.csv')

    employees = e.df2Dict()
    doc = Document()
    set_doc_style(doc)
    set_margins(doc)
    
    TimeSheet(
        doc,
        'Escola Municipal de Ensino Fundamental Dulcinéia Almeida do Nascimento',
        'FOLHA DE PONTO MENSAL DO (A) SERVIDOR (A)',
        employees[0]
    )
    doc.save('livro de ponto.docx')


