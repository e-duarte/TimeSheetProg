from docx import Document
from styles import set_header_doc_style, set_doc_style, set_margins, set_assign_style
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from time_sheet import TimeSheet
from calendar_util import CalendarUtil


HEADER_TEXT = """ESTADO DO PARÁ
PREFEITURA MUNICIPAL DE VITÓRIA DO XINGU
SECRETARIA MUNICIPAL DE EDUCAÇÃO
EMEF DULCINÉIA ALMEIDA DO NASCIMENTO
INEP 15111130"""

ASSIGN_TEXT_1 = """___________________________________
Luciana da Silva Gomes
Diretora Escolar
Port.Mun.0011/2022-SEMAD
"""

ASSIGN_TEXT_2 = """___________________________________
Danielle Juliana Cabral da Costa
Diretora Escolar
Port.Nº.0294/2022-SEMAD
EMEF. Dulcinéia Almeida do Nascimento
"""

class Book:
    def __init__(self, employees, month):
        self.doc = Document()
        self.employees = employees

        self.month = CalendarUtil().months.index(month) + 1
        print('month', self.month)
        
        set_doc_style(self.doc)
        set_margins(self.doc)
        
        self.build_header()
        
    
    def build_header(self,left_img = './images/bandeira-municipio.png', rigth_img = 'images/dulcineia.png'):
        header = self.doc.sections[0].header

        table = header.add_table(rows=1, cols=3, width=Cm(20.0))
    
        header_text_cell = table.rows[0].cells[1]
        header_text_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_header_text = header_text_cell.paragraphs[0].add_run(HEADER_TEXT)
        run_header_text.font.size = Pt(8)

        set_header_doc_style(table)
        
        #set images
        table.cell(0,0).paragraphs[0].add_run().add_picture(left_img, width=Cm(1.6))
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0,2).paragraphs[0].add_run().add_picture(rigth_img, width=Cm(1.6))
        table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def build_assigns(self):
        assign_table = self.doc.add_table(1,2)
        run = assign_table.cell(0,0).paragraphs[-1].add_run(ASSIGN_TEXT_1)
        run.font.bold = True
        run.font.size = Pt(8)
        
        run = assign_table.cell(0,1).paragraphs[-1].add_run(ASSIGN_TEXT_2)
        run.font.bold = True
        run.font.size = Pt(8)
        set_assign_style(assign_table)
    
    def add_line_metadata(self, par, line_data):
        formated_data = []
        for i, data in enumerate(line_data):
            if i == 0:
                formated_data.append((data[0], data[1]))
            else:
                formated_data.append((f'    {data[0]}', data[1]))

        for data in formated_data:
            #field
            run = par.add_run(f'{data[0]}')
            run.font.bold = True
            run.font.size = Pt(12)
            #value
            run = par.add_run(f'{data[1]}')
            run.font.size = Pt(12)
            
    def build_occurrence_page(self, employee):
        par = self.doc.add_paragraph()
        self.add_line_metadata(par, (('Nome completo do (a) Servidor (a): ', employee['Nome']),))
        self.add_line_metadata(par, 
            (
                ('\nEndereço: ', '_________________________'),
                ('Bairro: ', '_________________________'), 
            )
        )
        self.add_line_metadata(par, 
            (
                ('\nCelular: ', f'(93)_______________________' if employee['Telefone'] == '' else f'(93) {employee["Telefone"]}'),
            )
        )
        
        par = self.doc.add_paragraph()
        self.add_line_metadata(par, 
            (
                ('OCORRÊNCIAS: ', 'Data: _____/_____/_____.   Horário: _____:_____'),
            )
        )
        
        par = self.doc.add_paragraph()
        par.text = f'Encaminhamentos: {3*150*"_"}'
        
        par = self.doc.add_paragraph()
        par.text = f'Assinatura: {54*"_"}'
        
        par = self.doc.add_paragraph()
        self.add_line_metadata(par, (('ATESTADO: ', ''),))
        par = self.doc.add_paragraph()
        par.text = f'Data: _____/_____/_____.   Horário: _____:_____   Assinatura: {20*"_"}\n'
        par.text += f'Data: _____/_____/_____.   Horário: _____:_____   Assinatura: {20*"_"}\n'
        par.text += f'Data: _____/_____/_____.   Horário: _____:_____   Assinatura: {20*"_"}\n'
        par.text += f'Data: _____/_____/_____.   Horário: _____:_____   Assinatura: {20*"_"}'
        
        par = self.doc.add_paragraph()
        self.add_line_metadata(par, (('LICENÇA: ', ''),))
        par = self.doc.add_paragraph()
        par.text = f'Requerimento Nº {10*"_"}  Data: _____/_____/_____.\n'
        par.text += f'Portaria Nº {10*"_"}  Data: _____/_____/_____.\n'
        par.text += f'Período  de afastamento {44*"_"}\n'
        par.text += f'Assinatura: {54*"_"}'
        
        par = self.doc.add_paragraph()
        self.add_line_metadata(par, (('Férias: ', ''),))
        par = self.doc.add_paragraph()
        par.text = f'Requerimento Nº {10*"_"}  Data: _____/_____/_____.\n'
        par.text += f'Portaria Nº {10*"_"}  Data: _____/_____/_____.\n'
        par.text += f'Período  de afastamento {44*"_"}\n'
        par.text += f'Assinatura: {54*"_"}'
        
        par = self.doc.add_paragraph()
        par = self.doc.add_paragraph()
        # par = self.doc.add_paragraph()
        
        self.build_assigns()
        
    
    def build_timesheet_page(self, employee):
        institution = 'Escola Municipal de Ensino Fundamental Dulcinéia Almeida do Nascimento'
        title = 'FOLHA DE PONTO MENSAL DO (A) SERVIDOR (A)'
        TimeSheet(
            self.doc,
            institution,
            title,
            employee,
            self.month
        )
        
        self.doc.add_paragraph()
        
        self.build_assigns()
        
    def build_pages(self):
        for e in self.employees:
            print(f'Building Time Sheet page to {e["Nome"]}')
            self.build_timesheet_page(e)
            self.doc.add_page_break()
            self.build_occurrence_page(e)
            self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)




